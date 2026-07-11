from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_DIR = Path("reports/kona_biopsy_age_sensitivity")
DATA_PATH = OUT_DIR / "kona_biopsy_age_sensitivity_report_data.json"
DOCX_PATH = OUT_DIR / "Kona_Biopsy_Age_Model_Sensitivity_Report.docx"


def fmt(value, suffix: str = "") -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        if abs(value - round(value)) < 0.05:
            return f"{value:.0f}{suffix}"
        return f"{value:.1f}{suffix}"
    return f"{value}{suffix}"


def get_model(data: dict, key: str) -> dict:
    return next(item for item in data["summaries"] if item["key"] == key)


def get_eli(data: dict, key: str) -> dict:
    return next(item for item in data["eliComparisons"] if item["key"] == key)


def get_maturity(data: dict, model_key: str, variant: str) -> dict:
    return next(
        item
        for item in data["maturityAgeSensitivity"]
        if item["modelKey"] == model_key and item["maturityVariant"] == variant
    )


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def model_short(label: str) -> str:
    return (
        label.replace("Model 1 + ", "")
        .replace("Model 1: ", "Model 1: ")
        .replace("Best evidence model: ", "Best evidence: ")
        .replace(" with maturity-size/age assumptions", "")
        .replace(" with maturity-age assumptions", "")
        .replace(" dated/first-sighting", "")
    )


def build_doc():
    data = json.loads(DATA_PATH.read_text())
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_heading("Kona Biopsy Age Model Sensitivity Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        f"Kona biopsied manta rays (n={data['recordCount']}); age estimates as of {data['ageReferenceDate']}; "
        f"generated from the Hawaii Manta Tracker age-ranking workbench."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "This analysis compares five candidate evidence models for estimating minimum plausible age, rank order, and "
        "pairwise generation plausibility among Kona biopsied manta rays. Model 1, earliest known first sighting, is "
        "used as the baseline because it is direct observation and is least dependent on interpretation. The recommended "
        "parsimonious model for internal analysis is Models 1-4: first sighting, pup-at-first-sighting, HAMER size evidence "
        "with maturity-size assumptions, and HAMER dated age-class evidence with maturity-age assumptions, excluding MPRF "
        "age-class labels until those labels can be independently validated."
    )
    best = get_model(data, "best_evidence_no_mprf_class")
    all_models = get_model(data, "all_models")
    m5 = get_model(data, "m5_mprf_age_class_assumption")
    add_body(
        doc,
        f"The recommended evidence model changed minimum age for only {best['minAge']['changedCount']} of 84 animals, "
        f"with a maximum increase of {fmt(best['minAge']['maxIncrease'])} years and only "
        f"{best['ranks']['relativeOrderChangedCount']} relative-order rank changes. It therefore preserves the robust "
        "first-sighting ranking structure while adding bounded biological evidence where HAMER data are available. "
        f"In contrast, adding MPRF age-class labels increased ages for {m5['minAge']['changedCount']} animals, caused "
        f"{m5['ranks']['changedAbsoluteCount']} absolute rank changes, and reduced discrepancy with Eli's size-adjusted "
        "ranking, but this effect is concentrated in a few records whose MPRF age labels need validation."
    )
    add_body(
        doc,
        "Pairwise generation diagnostics should be interpreted as a screening tool for plausible generation separation, "
        "not as parentage inference. The matrix flags whether one animal could be a parent (P), child (C), same generation "
        "(S), or unresolved (U) based on actual dated adult-versus-juvenile evidence. Back-calculated age intervals are used "
        "for ranking, but they are not used to back-date adult status for P/C calls. Under the recommended Models 1-4 midpoint setting, 45 mantas had at least "
        "one plausible parent-position call, 82 had at least one plausible child-position call, and the matrix produced 1,173 "
        "directional P-C relationships. These are candidate generation-gap screens, not genetic parentage assignments. Genetic "
        "data are still required to confirm parent-offspring or sibling relationships."
    )
    best_low = get_maturity(data, "best_evidence_no_mprf_class", "low")
    best_mid = get_maturity(data, "best_evidence_no_mprf_class", "midpoint")
    best_high = get_maturity(data, "best_evidence_no_mprf_class", "high")
    add_body(
        doc,
        f"Sensitivity testing across the evaluated maturity-age range showed that age ranking was stable but generation calls were "
        f"threshold-sensitive. Under the recommended model, changing from low maturity ages (male 5, female 8) to high maturity ages "
        f"(male 8, female 15) shifted the number of mantas with at least one P call from {best_low['mantasWithP']} to "
        f"{best_high['mantasWithP']} and directional P-C relationships from {best_low['pcRelationships']} to "
        f"{best_high['pcRelationships']}, while relative-order rank changes remained modest ({best_high['relativeOrderChangedVsMidpoint']} "
        f"under the high setting relative to the midpoint default)."
    )

    add_heading(doc, "Email-ready summary for Eli", 1)
    add_body(
        doc,
        "Hi Eli - I assessed five age-evidence models for the Kona biopsy animals and compared each one against the baseline model, "
        "which uses only earliest known first sighting to estimate minimum age as of 1/1/2024. I used first sighting as the baseline "
        "because it is direct observation and does not require age-class, size, or maturity assumptions. The other models added: "
        "pup-at-first-sighting labels; HAMER size measurements interpreted with sex-specific maturity-size and maturity-age assumptions; "
        "HAMER dated age-class determinations interpreted with sex-specific maturity ages; and MPRF age-class labels interpreted the same way."
    )
    add_body(
        doc,
        "The main result is that the HAMER-based additions did not substantially move the needle on minimum age estimates or rank order. "
        "Models 2-4 caused only small changes relative to the first-sighting baseline, and even when I shifted maturity ages across the "
        "evaluated low-to-high range (males 5-8 years; females 8-15 years), rank order remained fairly stable. Those maturity-age assumptions "
        "did matter more for the pairwise generation screen because the P/C/S/U matrix depends directly on whether the age separation clears "
        "the selected maturity threshold. Under the recommended HAMER-prioritized model, changing from the low to high maturity setting reduced "
        "directional P-C relationships from 1,573 to 837 and reduced mantas with at least one P call from 56 to 35, while minimum-age ranks "
        "changed only modestly."
    )
    add_body(
        doc,
        "A few examples show how the models behave. Ho Ray is affected by HAMER size evidence, increasing from 16.2 to 17.8 years under "
        "the midpoint maturity setting. Otho D Ray is affected by HAMER age class, increasing from 8.4 to 9.9 years. Zach a ray changes only "
        "modestly, from 19.2 to 19.6 years. The MPRF age-class model moved the needle more for a few animals, especially Lefty/MP001, Big "
        "Bertha/MP009, Mango/MP138, Celestial Ray/MP275, Takahashi Ray/MP144, and Tim/MP197, because adult or mature labels at first sighting "
        "add the selected sex-specific maturity age. That MPRF-inclusive model was the closest match to your size-adjusted age rankings, which "
        "suggests your approach is most similar to my Model 1 + Model 5 combination."
    )
    add_body(
        doc,
        "One identifier issue still needs to be resolved before attaching sequence data to individual animals. In the HAMER source data, "
        "Lefty is biopsy code BK6, pk_biopsy_id 57, catalog 704, and Independence Ray is BK7, pk_biopsy_id 58, catalog 705; both were sampled "
        "on 2012-09-03. In the current imported database, PSN24OQ691 / BI_K26 is stored on Independence Ray, while Lefty's final biopsy row has "
        "no stored lab/sample ID. However, the Jonathan/Whitney staging row for PSN24OQ691 / BI_K26 contains the HAMER text 'BK6 ... MP1 Lefty "
        "Cat362 F BK6' and also contains mixed Lefty/MP184 wording. BI_K27 appears in that staging table for Independence/BK7 but is not present "
        "in the final biopsy table. This looks like a crossover/crosswalk issue between the adjacent BK6 and BK7 records rather than a problem "
        "with the original HAMER biopsy records. It does not change the age-sensitivity conclusions, but it should be resolved before using that "
        "sequence assignment in parentage comparisons."
    )

    add_heading(doc, "Introduction", 1)
    add_body(
        doc,
        "The purpose of this analysis is to determine which available evidence sources materially improve age estimates and "
        "rank ordering for Kona biopsied reef manta rays while avoiding unnecessary assumptions. Age estimates are needed for "
        "close-kin and parentage interpretation because biologically plausible parent-offspring or sibling hypotheses depend on "
        "minimum age, maturity age, and generation separation. Published work on Hawaiian manta rays supports sex-specific "
        "maturity thresholds and reproductive maturity assumptions, while close-kin demographic work emphasizes explicit treatment "
        "of aging error and uncertainty (Deakos 2010, 2011, 2012; Swenson et al. 2024)."
    )

    add_heading(doc, "Methods", 1)
    add_body(
        doc,
        "All models estimated minimum plausible age as of 1/1/2024. Records were ranked from oldest to youngest by final minimum "
        "plausible age, with ties preserved and stable ordering based on existing rank and biopsy identifier. Model 1 used only the "
        "earliest known first-sighting date. Model 2 added direct birth-anchor evidence from first sightings labelled as pup. Model 3 "
        "added dated HAMER size measurements interpreted using maturity-size thresholds of 2.8 m for males and 3.37 m for females and "
        "maturity ages of 6.5 and 11.5 years, respectively. Model 4 added HAMER dated age-class observations interpreted with the same "
        "sex-specific maturity ages. Model 5 added MPRF dated or first-sighting age-class labels interpreted with the same maturity ages. "
        "The primary evidence model recommended here combines Models 1-4 and excludes MPRF age class until those labels are validated."
    )
    add_body(
        doc,
        "Concrete examples help show how the models behave. Ho Ray illustrates Model 3: first sighting alone estimated 16.2 years, "
        "but a dated HAMER size measurement interpreted with maturity-size and maturity-age assumptions raised the minimum age to "
        "17.8 years. Otho D Ray illustrates Model 4: HAMER age-class evidence at biopsy raised the estimate from 8.4 to 9.9 years. "
        "Zach a ray was affected by both HAMER age class and size evidence, but only modestly, increasing from 19.2 to 19.6 years. "
        "By contrast, Lefty and Big Bertha illustrate why Model 5 is treated separately: MPRF adult labels at first sighting add the "
        "female maturity-age assumption, raising Lefty from 44.6 to 56.1 years and Big Bertha from 32.0 to 43.5 years under the "
        "midpoint maturity setting."
    )
    add_body(
        doc,
        "Growth-rate extrapolation was not used for primary age ranking or pairwise generation diagnostics. Available growth-rate evidence "
        "for this population remains too sparse and stage-dependent to support a reliable linear age assignment, and estimates from captive "
        "or non-local contexts may not transfer cleanly to wild Hawaiian reef mantas. Growth-rate fields are therefore retained only as "
        "exploratory sensitivity inputs for future size/birth-year work, not as part of the current primary evidence model."
    )
    add_body(
        doc,
        "Pairwise generation diagnostics classify each ordered pair as P, C, S, or U. P means the row animal could be a generational "
        "parent of the column animal under the selected maturity-age criteria; C is the reciprocal child state; S means available "
        "bounds support same-generation classification rather than a maturity-age gap; U means unresolved. The diagnostic uses dated "
        "adult-versus-juvenile evidence only. If dated adult-versus-juvenile evidence is unavailable, the pair remains unresolved rather "
        "than using back-calculated age intervals to infer an earlier adult date."
    )
    add_body(
        doc,
        "Comparisons to Eli's spreadsheet used the size-adjusted age estimate and age rank from the 'sample list' sheet. Matching was "
        "by Sample_ID when consistent and by MPRF catalog ID when Sample_ID was absent. A Sample_ID/MPRF conflict involving BI_K26 was "
        "flagged and not allowed to distort the comparison. A live database review showed that the final biopsy table stores "
        "PSN24OQ691 / BI_K26 on Independence Ray (pk_biopsy_id 58), while Lefty's final biopsy row (pk_biopsy_id 57) has no stored "
        "lab/sample ID. The Jonathan/Whitney staging row for PSN24OQ691 / BI_K26 contains HAMER BK6/Lefty text but also mixed MP184 wording, "
        "and BI_K27 appears in staging for Independence/BK7 but is absent from the final biopsy table. This pattern is consistent with an "
        "external sample crosswalk issue between adjacent same-day records, not an inconsistency in the original HAMER BK6/BK7 biopsy records."
    )

    add_heading(doc, "Results", 1)
    add_body(doc, "Table 1 summarizes the age and rank sensitivity of each model relative to Model 1.")
    model_rows = []
    for key in ["m1", "m2_pup", "m3_size_assumption", "m4_hamer_age_class_assumption", "m5_mprf_age_class_assumption", "best_evidence_no_mprf_class", "all_models"]:
        s = get_model(data, key)
        model_rows.append([
            model_short(s["label"]),
            fmt(s["minAge"]["mean"]),
            fmt(s["minAge"]["changedCount"]),
            fmt(s["minAge"]["maxIncrease"]),
            fmt(s["ranks"]["changedAbsoluteCount"]),
            fmt(s["ranks"]["relativeOrderChangedCount"]),
            fmt(s["ranks"]["maxAbsRankDelta"]),
        ])
    add_table(
        doc,
        ["Model", "Mean min age", "Age changed", "Max age increase", "Abs. rank changes", "Relative-order changes", "Max rank delta"],
        model_rows,
    )

    add_body(
        doc,
        "The first-sighting baseline produced a mean minimum age of 12.6 years. Pup-at-first-sighting evidence did not alter minimum "
        "age or rank because the first sighting already establishes the minimum-age floor, but it strongly improved upper-bound and "
        "same-generation interpretation. HAMER size and HAMER age-class evidence each changed only two animals when interpreted with "
        "maturity assumptions. The largest individual increases in the recommended evidence model were Ho Ray (+1.7 years), Otho D Ray "
        "(+1.5 years), and Zach a ray (+0.4 years)."
    )

    top_age = best["topAgeIncreases"]
    add_table(
        doc,
        ["Animal", "Evidence model age", "Baseline age", "Increase", "Baseline rank", "Model rank", "Rank delta"],
        [[r["name"], fmt(r["scenarioMinAge"]), fmt(r["baselineMinAge"]), fmt(r["ageIncrease"]), fmt(r["baselineRank"]), fmt(r["scenarioRank"]), fmt(r["rankDelta"])] for r in top_age],
    )

    add_body(doc, "Table 3 summarizes pairwise generation diagnostics by model.")
    pairwise_rows = []
    for key in ["m1", "m2_pup", "m3_size_assumption", "m4_hamer_age_class_assumption", "m5_mprf_age_class_assumption", "best_evidence_no_mprf_class", "all_models"]:
        s = get_model(data, key)
        pairwise_rows.append([
            model_short(s["label"]),
            fmt(s["pairwise"]["animalsWithP"]),
            fmt(s["pairwise"]["animalsWithC"]),
            fmt(s["pairwise"]["animalsWithS"]),
            fmt(s["pairwise"]["animalsWithU"]),
            fmt(s["pairwise"]["pcRelationships"]),
            fmt(s["pairwise"]["sameGenerationRelationships"]),
        ])
    add_table(
        doc,
        ["Model", "Mantas with P", "Mantas with C", "Mantas with S", "Mantas with U", "P-C relationships", "S relationships"],
        pairwise_rows,
    )
    add_body(
        doc,
        "Pup evidence had the largest effect on same-generation resolution: S relationships increased from 450 under Model 1 to 1,207 "
        "under Model 2, while P-C relationship counts remained unchanged. This means pup evidence improves exclusion/resolution without "
        "reshuffling minimum-age rank. The recommended evidence model retained nearly the same parent-child screening count as Model 1 "
        "(1,173 vs. 1,179 P-C relationships) while increasing S relationships to 1,236."
    )

    add_heading(doc, "Maturity-age Sensitivity", 2)
    add_body(
        doc,
        "Because generation separation depends directly on the selected age at maturity, the recommended model and the all-models "
        "MPRF-inclusive model were rerun at the lower and upper maturity-age ranges. The low setting used male 5 years and female 8 "
        "years. The midpoint/default setting used male 6.5 years and female 11.5 years. The high setting used male 8 years and female "
        "15 years."
    )
    maturity_rows = []
    maturity_labels = {
        "best_evidence_no_mprf_class": "Recommended Models 1-4",
        "all_models": "All models incl. MPRF",
    }
    for model_key in ["best_evidence_no_mprf_class", "all_models"]:
        for variant in ["low", "midpoint", "high"]:
            r = get_maturity(data, model_key, variant)
            maturity_rows.append([
                maturity_labels[model_key],
                f"{fmt(r['maleMaturityAgeYears'])}/{fmt(r['femaleMaturityAgeYears'])}",
                fmt(r["meanMinimumAge"]),
                fmt(r["ageChangedVsMidpoint"]),
                fmt(r["relativeOrderChangedVsMidpoint"]),
                fmt(r["mantasWithP"]),
                fmt(r["mantasWithC"]),
                fmt(r["pcRelationships"]),
                fmt(r["sameGenerationRelationships"]),
                fmt(r["cellU"]),
            ])
    add_table(
        doc,
        ["Model", "M/F maturity age", "Mean min age", "Age changes", "Relative rank changes", "Mantas P", "Mantas C", "P-C relationships", "S relationships", "U cells"],
        maturity_rows,
    )
    add_body(
        doc,
        "This sensitivity test confirmed the expected pattern. Maturity-age assumptions had little influence on minimum age and rank "
        "because relatively few animals have dated adult/juvenile anchors that add maturity years. However, maturity-age assumptions "
        "strongly changed the pairwise generation matrix. Lower maturity ages made it easier for older animals to clear the generation "
        "gap threshold, increasing P-C relationships. Higher maturity ages made the test stricter, reducing P-C relationships and "
        "increasing same-generation or unresolved classifications."
    )
    add_body(
        doc,
        "Real examples make this visible. Under the recommended model, Ho Ray's minimum age was 17.8 years at the midpoint setting; "
        "it dropped to 16.3 years under the low setting and rose to 19.3 years under the high setting because his size-derived maturity "
        "interpretation is tied to male maturity age. Otho D Ray showed the same pattern through HAMER age-class evidence: 8.4 years "
        "at the low setting, 9.9 years at the midpoint, and 11.4 years at the high setting. When MPRF age class is included, the effect "
        "is larger for adult-at-first-sighting females: Lefty shifts from 52.6 years at the low setting to 56.1 years at midpoint and "
        "59.6 years at high; Big Bertha shifts from 40.0 to 43.5 to 47.0 years; Mango shifts from 23.6 to 27.1 to 30.6 years. These "
        "examples show why maturity-age sensitivity is essential for generation diagnostics even when the rank order remains fairly stable."
    )

    add_body(doc, "Table 4 lists the strongest plausible parent contributors under the recommended evidence model.")
    add_table(
        doc,
        ["Rank", "Animal", "Sample_ID", "HAMER catalog", "MPRF catalog", "Min age", "P", "C", "S", "U"],
        [[r["rank"], r["name"], r.get("sampleId") or "-", r["hamerCatalogId"], r["mprfCatalogId"], fmt(r["minimumAge"]), r["P"], r["C"], r["S"], r["U"]] for r in data["topParentContributors"][:10]],
    )

    add_heading(doc, "Comparison With Eli's Size-adjusted Estimates", 1)
    add_body(
        doc,
        "The model including MPRF age-class labels aligned most closely with Eli's size-adjusted ages, reducing the mean absolute age "
        "difference from 2.0 years under Model 1 to 1.6 years and reducing records differing by five or more years from 8 to 3. This is "
        "expected because Eli's table appears to incorporate MPRF age-class or size-adjusted interpretation that adds maturity-age years "
        "to some animals classified as adult at or near first sighting."
    )
    eli_rows = []
    for key in ["m1", "m2_pup", "m3_size_assumption", "m4_hamer_age_class_assumption", "m5_mprf_age_class_assumption", "best_evidence_no_mprf_class", "all_models"]:
        e = get_eli(data, key)
        eli_rows.append([
            model_short(e["label"]),
            e["matched"],
            e["unmatchedWorkbench"],
            fmt(e["meanAgeDifference"]),
            fmt(e["meanAbsAgeDifference"]),
            fmt(e["recordsAgeDiffGte5"]),
            fmt(e["meanAbsRankDifference"]),
            fmt(e["recordsRankDiffGte10"]),
        ])
    add_table(
        doc,
        ["Model", "Matched", "Unmatched", "Mean age diff", "Mean abs age diff", "|Age diff| >= 5", "Mean abs rank diff", "|Rank diff| >= 10"],
        eli_rows,
    )
    best_eli = get_eli(data, "best_evidence_no_mprf_class")
    add_table(
        doc,
        ["Workbench animal", "Eli animal", "Sample_ID", "MPRF", "Workbench age", "Eli size-adjusted age", "Age diff", "Workbench rank", "Eli rank", "Rank diff"],
        [
            [
                r["name"],
                r.get("eliName") or "-",
                r.get("sampleId") or r.get("eliSampleId") or "-",
                r.get("mprfCatalogId") or r.get("eliMprfCatalogId") or "-",
                fmt(r.get("workbenchMinimumAge")),
                fmt(r.get("eliSizeAdjustedAge")),
                fmt(r.get("ageDifferenceYears")),
                fmt(r.get("workbenchRank")),
                fmt(r.get("eliAgeRank")),
                fmt(r.get("rankDifference")),
            ]
            for r in best_eli["topAgeDifferences"][:8]
        ],
    )

    add_heading(doc, "Discussion", 1)
    add_body(
        doc,
        "The best balance of parsimony and biological value is the HAMER-prioritized evidence model: Models 1-4, excluding MPRF age "
        "class. It is conservative, reproducible, and minimally changes age rankings while adding useful birth-anchor, size, and dated "
        "age-class information. It also avoids over-weighting unvalidated MPRF age-class labels. This model should be the default for "
        "internal ranking and sensitivity reporting."
    )
    add_body(
        doc,
        "Model 5 should be retained as a comparison model rather than discarded. It answers a critical sensitivity question: what changes "
        "when MPRF age-class labels are allowed into the model? The answer is that the mean age changes only modestly, but rank changes "
        "are concentrated and sometimes large. Lefty, Big Bertha, Mango, Celestial Ray, Takahashi Ray, and Tim are the primary movers "
        "because adult or mature MPRF labels at first sighting add the selected maturity age to the first-sighting minimum."
    )
    add_body(
        doc,
        "The pairwise matrix is useful for screening close-kin hypotheses because it can identify pairs that are too close in apparent age "
        "to be plausible parent-offspring pairs and pairs with enough separation to remain plausible. However, high P totals should not be "
        "interpreted as evidence of many offspring. They indicate that older, well-documented animals such as Lefty, Hook, Scar, Big Bertha, "
        "Curly Ray, and Sugar Ray remain plausible generation-older candidates for many younger animals under age-only rules."
    )
    add_body(
        doc,
        "The comparison with Eli's results suggests two practical next steps. First, resolve the BI_K26/BI_K27 crossover between "
        "Lefty/BK6 and Independence Ray/BK7 before using either record in formal sequence-linked comparisons. Because Lefty and Independence "
        "are different sexes, any available sequence-confirmed sex call may help verify which sample belongs to which animal. Second, review "
        "the small set of large discrepancies to determine whether Eli's "
        "size-adjusted estimates are using additional size-growth assumptions, MPRF age-class assumptions, or catalog-history evidence that "
        "should be explicitly represented in the workbench."
    )

    add_heading(doc, "Supplemental Materials", 1)
    add_body(
        doc,
        "Supplemental workbook: kona_biopsy_age_sensitivity_tables.xlsx. Supplemental JSON: "
        "kona_biopsy_age_sensitivity_report_data.json. These files contain the complete model summaries, Eli comparisons, and top pairwise "
        "contributors used to generate this report."
    )

    add_heading(doc, "Prompt for External Polishing", 1)
    add_body(
        doc,
        "Use this prompt only if a separate writing tool is used: 'Polish the attached Kona biopsy age-model sensitivity report without "
        "changing the numerical results. Preserve all model definitions, the recommendation that Models 1-4 are the parsimonious HAMER-prioritized "
        "evidence model, and the caution that Model 5/MPRF age class is a sensitivity model pending validation. Improve flow, add concise "
        "scientific transitions, and keep claims conservative. Do not convert pairwise P/C/S/U calls into genetic parentage conclusions.'"
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
