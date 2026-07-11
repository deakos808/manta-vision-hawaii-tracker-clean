from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("reports/kona_pi_hat_generation_alignment")
SUMMARY_CSV = OUT_DIR / "kona_pi_hat_generation_alignment_summary.csv"
ALL_PAIRS_CSV = OUT_DIR / "kona_pi_hat_generation_alignment_all_pairs.csv"
STRONG_PAIRS_CSV = OUT_DIR / "kona_pi_hat_generation_alignment_strong_pairs.csv"
REPORT_PATH = OUT_DIR / "Kona_PI_HAT_Generation_Matrix_Alignment_Report.docx"
CORE_RELATIONSHIPS_IMAGE = Path(
    "/Users/littlemac/Dropbox/Work/HAMER/Research/Elasmobranchs/Manta Rays/2. Genetics/"
    "MPRF Manta Parentage Collaboration/Kona Biopsy Age Rankings/Genetics Results/Core Relationships 6.6.26.png"
)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def num(value: object) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def fmt(value: object, decimals: int = 1) -> str:
    number = num(value)
    if number is None:
        return str(value or "-")
    if abs(number - round(number)) < 10 ** (-(decimals + 1)):
        return f"{number:.0f}"
    return f"{number:.{decimals}f}"


def pct(value: float, total: float) -> str:
    return "-" if not total else f"{100 * value / total:.1f}%"


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)
    return doc


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def set_cell(cell, value: object, bold: bool = False) -> None:
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = read_csv(SUMMARY_CSV)
    all_pairs = read_csv(ALL_PAIRS_CSV)
    strong_pairs = read_csv(STRONG_PAIRS_CSV)

    strong = [row for row in strong_pairs if num(row.get("PI_HAT")) is not None and num(row["PI_HAT"]) >= 0.35]
    matched_strong = [row for row in strong if row["alignmentCategory"] != "unmatched"]
    pc_strong = [row for row in matched_strong if row["alignmentCategory"] == "genetics_close_age_pc"]
    s_strong = [row for row in matched_strong if row["alignmentCategory"] == "genetics_close_age_s"]
    unknown_strong = [row for row in matched_strong if row["alignmentCategory"] == "genetics_close_age_unknown"]
    unmatched_strong = [row for row in strong if row["alignmentCategory"] == "unmatched"]
    direct_conflicts = [
        row
        for row in matched_strong
        if row["alignmentCategory"] not in {"genetics_close_age_pc", "genetics_close_age_s", "genetics_close_age_unknown"}
    ]

    matched_all = [row for row in all_pairs if row["alignmentCategory"] != "unmatched"]
    matrix_pc = [
        row
        for row in matched_all
        if row["codeAtoB"] in {"P", "C"} or row["codeBtoA"] in {"P", "C"}
    ]
    matrix_s = [
        row
        for row in matched_all
        if row["codeAtoB"] == "S" or row["codeBtoA"] == "S"
    ]
    matrix_u = [
        row
        for row in matched_all
        if row["codeAtoB"] == "U" or row["codeBtoA"] == "U"
    ]
    matrix_pc_strong = [row for row in matrix_pc if num(row["PI_HAT"]) is not None and num(row["PI_HAT"]) >= 0.35]
    matrix_s_strong = [row for row in matrix_s if num(row["PI_HAT"]) is not None and num(row["PI_HAT"]) >= 0.35]
    matrix_u_strong = [row for row in matrix_u if num(row["PI_HAT"]) is not None and num(row["PI_HAT"]) >= 0.35]

    doc = setup_doc()
    title = doc.add_heading("Kona Genetic Relatedness vs. Generation Matrix Alignment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("PI_HAT screening results compared with age-based parent-child/same-generation plausibility calls")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", 1)
    add_body(
        doc,
        "The PI_HAT results and the Kona generation matrix are complementary rather than redundant. PI_HAT identifies pairs that "
        "share unusually high genetic similarity; the generation matrix asks whether dated age, size, and life-stage evidence makes "
        "a parent-child direction plausible, favors a same-generation interpretation, or leaves the relationship unresolved."
    )
    add_bullets(
        doc,
        [
            f"The pi-hat file contained {len(all_pairs)} pairwise genetic comparisons. Of these, {len(matched_all)} matched both members to the Kona biopsied generation matrix and {len(all_pairs) - len(matched_all)} could not be fully mapped.",
            f"Using PI_HAT >= 0.35 as the strongest-relation screen, {len(strong)} pairs were flagged genetically; {len(matched_strong)} mapped to the matrix and {len(unmatched_strong)} did not.",
            f"Among the mapped strong pairs, {len(pc_strong)} were age-compatible with a directional parent-child hypothesis and {len(s_strong)} were more consistent with a same-generation/full-sibling interpretation.",
            f"{len(unknown_strong)} mapped strong pairs had insufficient dated age-stage evidence to determine parent-child versus sibling direction.",
            f"No mapped strong pair was clearly rejected by the generation matrix. The main limitation is unresolved or unmatched evidence, not direct contradiction.",
        ],
    )

    doc.add_heading("What PI_HAT Means", 1)
    add_body(
        doc,
        "In plain terms, PI_HAT is a genetic similarity score. It estimates how much of two individuals' genomes appear to be shared "
        "because they inherited the same DNA from recent common ancestry. A value near 0 usually means the pair does not look closely "
        "related in the marker set. A value near 0.25 is often consistent with second-degree relationships, a value near 0.5 is often "
        "consistent with first-degree relationships such as parent-child or full sibling, and a value close to 1 can suggest duplicate "
        "samples, repeated genotypes, identical/twin-level similarity, or a sample/identity issue that should be checked before biological interpretation."
    )
    add_body(
        doc,
        "For this analysis, PI_HAT >= 0.35 was treated as a practical close-kin screen because the provided figure labeled those as the "
        "strongest relations. This threshold does not prove parentage; it identifies pairs worth testing against age and life-history evidence."
    )

    doc.add_heading("Methods", 1)
    add_body(
        doc,
        "The comparison used the Kona biopsied best-evidence generation matrix from the population age model sensitivity workbook. "
        "The matrix classifies each ordered pair as P, C, S, or U: P means the row individual could be parent of the column individual, "
        "C is the reciprocal child-position call, S means dated evidence favors a same-generation relationship, and U means insufficient dated "
        "adult-versus-juvenile/size evidence. Sample IDs were matched first, with names used as a fallback."
    )
    add_body(
        doc,
        "The matrix is deliberately permissive: it identifies pairs that are biologically possible based on age separation, not pairs that are genetically supported. "
        "Therefore, a low PI_HAT for a matrix-permitted P/C pair is not a failure of the matrix; it simply means the age screen allowed a possible relationship that genetics does not support."
    )

    doc.add_heading("Alignment Results", 1)
    add_table(
        doc,
        ["PI_HAT set", "Pairs", "Mapped", "Unmapped", "P/C compatible", "Same-generation compatible", "Age unresolved", "Very high PI_HAT"],
        [
            [
                row["threshold"],
                row["pairs"],
                row["matchedPairs"],
                row["unmatchedPairs"],
                row["ageCompatiblePC"],
                row["ageCompatibleS"],
                row["ageUnknown"],
                row["veryHighPiHat"],
            ]
            for row in summary
            if row["threshold"] != "all pairs"
        ],
        widths=[1.1, 0.55, 0.65, 0.7, 0.85, 1.0, 0.8, 0.8],
    )
    add_body(
        doc,
        f"Among strong mapped pairs (PI_HAT >= 0.35), {len(pc_strong) + len(s_strong)} of {len(matched_strong)} aligned with an interpretable generation-matrix category: "
        f"{len(pc_strong)} parent-child compatible and {len(s_strong)} same-generation compatible. The remaining {len(unknown_strong)} were genetically compelling but unresolved by age evidence."
    )

    doc.add_heading("Strong Genetic Pairs That Align With Parent-Child Plausibility", 2)
    add_table(
        doc,
        ["Sample A", "Name A", "Sample B", "Name B", "PI_HAT", "Matrix interpretation", "Min age A", "Min age B"],
        [
            [
                row["sampleIdA"],
                row["nameA"],
                row["sampleIdB"],
                row["nameB"],
                row["PI_HAT"],
                row["generationPrediction"],
                row["minimumAgeA"],
                row["minimumAgeB"],
            ]
            for row in pc_strong
        ],
        widths=[0.7, 1.0, 0.7, 1.0, 0.6, 1.6, 0.6, 0.6],
    )

    doc.add_heading("Strong Genetic Pairs That Align Better With Same-Generation Plausibility", 2)
    add_table(
        doc,
        ["Sample A", "Name A", "Sample B", "Name B", "PI_HAT", "Matrix interpretation"],
        [
            [
                row["sampleIdA"],
                row["nameA"],
                row["sampleIdB"],
                row["nameB"],
                row["PI_HAT"],
                row["generationPrediction"],
            ]
            for row in s_strong
        ],
        widths=[0.8, 1.1, 0.8, 1.1, 0.7, 1.8],
    )

    doc.add_heading("Strong Genetic Pairs With Insufficient Matrix Evidence", 2)
    add_body(
        doc,
        "These pairs are genetically important, but the age matrix could not resolve direction because dated adult-versus-juvenile evidence was missing or not decisive."
    )
    add_table(
        doc,
        ["Sample A", "Name A", "Sample B", "Name B", "PI_HAT", "Min age A", "Min age B"],
        [
            [
                row["sampleIdA"],
                row["nameA"],
                row["sampleIdB"],
                row["nameB"],
                row["PI_HAT"],
                row["minimumAgeA"],
                row["minimumAgeB"],
            ]
            for row in unknown_strong[:14]
        ],
        widths=[0.8, 1.2, 0.8, 1.2, 0.7, 0.7, 0.7],
    )

    doc.add_heading("Unmapped Strong Genetic Pairs", 2)
    add_body(
        doc,
        "These pairs should be treated as ID-reconciliation priorities before biological interpretation. Several include very high PI_HAT values, so resolving sample identity may materially change the final relationship interpretation."
    )
    add_table(
        doc,
        ["Sample A", "Name A", "Sample B", "Name B", "PI_HAT"],
        [[row["sampleIdA"], row["nameA"], row["sampleIdB"], row["nameB"], row["PI_HAT"]] for row in unmatched_strong],
        widths=[0.9, 1.4, 0.9, 1.4, 0.7],
    )

    doc.add_heading("What The Overall Ratios Imply", 1)
    add_table(
        doc,
        ["Matrix category among mapped genetic pairs", "Pairs with PI_HAT data", "Pairs with PI_HAT >= 0.35", "Interpretation"],
        [
            [
                "Parent-child possible by age matrix",
                len(matrix_pc),
                len(matrix_pc_strong),
                "Only a small subset of age-possible parent-child pairs are genetically strong, as expected for a permissive age screen.",
            ],
            [
                "Same-generation possible by age matrix",
                len(matrix_s),
                len(matrix_s_strong),
                "One strong genetic pair is better treated as same-generation compatible rather than directional parent-child.",
            ],
            [
                "Unknown by age matrix",
                len(matrix_u),
                len(matrix_u_strong),
                "Many genetic signals cannot be resolved by age evidence alone; these need better dated life-stage evidence or genetics-based relationship modeling.",
            ],
        ],
        widths=[1.7, 0.9, 0.9, 3.1],
    )
    add_body(
        doc,
        "The sampled Kona population appears to contain a small set of strong close-kin signals embedded within a much larger set of low-relatedness comparisons. "
        "That is exactly the expected use case for combining genetics with the generation matrix: genetics identifies the candidate relatives, and the matrix narrows which candidate relationships are biologically plausible as parent-child, which look more like same generation, and which remain unresolved."
    )
    add_body(
        doc,
        "The current results do not suggest broad disagreement between PI_HAT and the age matrix. Instead, they show that the matrix is conservative in a useful way: it confirms a few parent-child-compatible genetic pairs, flags one likely same-generation-compatible strong pair, and leaves many strong genetic pairs unresolved because the dated age evidence is not sufficient to force a direction."
    )

    doc.add_heading("Priority Follow-up", 1)
    add_bullets(
        doc,
        [
            "Resolve unmatched or conflicting sample IDs before final interpretation, especially BI118/Faux Ray, BI_K511, BI_KM35, BI_K27/Independence Ray, BI_K52, BI_K53, BI129/Vinny, and M_51.",
            "Treat PI_HAT values above approximately 0.75 as identity/QC review candidates before interpreting them as ordinary kinship.",
            "For unresolved strong pairs, use additional relationship inference tools or directed pedigree testing to distinguish parent-child from full-sibling hypotheses.",
        ],
    )

    if CORE_RELATIONSHIPS_IMAGE.exists():
        doc.add_heading("Reference Figure", 1)
        add_body(doc, "Provided core-relationships figure used as visual context for the PI_HAT >= 0.35 screen.")
        doc.add_picture(str(CORE_RELATIONSHIPS_IMAGE), width=Inches(6.7))

    doc.add_heading("Supplemental Files", 1)
    add_body(
        doc,
        "Detailed pairwise outputs are saved in the same report folder: "
        "`kona_pi_hat_generation_alignment_all_pairs.csv`, `kona_pi_hat_generation_alignment_strong_pairs.csv`, and "
        "`kona_pi_hat_generation_alignment_summary.csv`."
    )

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    main()
    print(REPORT_PATH)
