import argparse
import csv
import os
from collections import Counter
from datetime import datetime

import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


BASE = None
OUT_DIR = None
FILES = {
    "model1": {
        "filename": "kona_model1_high_confidence_age_rank.xlsx",
        "rank_sheet": "Model 1 Ranks",
        "evidence_sheet": "Evidence by Source",
        "rank_col": "model1_rank",
        "score_col": "model1_score",
        "evidence_col": "model1_evidence",
        "basis_col": "model1_evidence_basis",
    },
    "model2": {
        "filename": "kona_model2_mprf_age_classes_age_rank.xlsx",
        "rank_sheet": "Model 2 Ranks",
        "evidence_sheet": "Evidence by Source",
        "rank_col": "model2_rank",
        "score_col": "model2_score",
        "evidence_col": "model2_evidence",
        "basis_col": "model2_evidence_basis",
    },
    "model3": {
        "filename": "kona_model3_pup_birth_year_age_rank.xlsx",
        "rank_sheet": "Model 3 Ranks",
        "evidence_sheet": "Evidence by Source",
        "rank_col": "model3_rank",
        "score_col": "model3_score",
        "evidence_col": "model3_evidence",
        "basis_col": "model3_evidence_basis",
    },
}


def configure_paths(input_dir, output_dir):
    """Configure paths that were machine-specific in the May 29, 2026 script."""
    global BASE, OUT_DIR
    BASE = os.path.abspath(input_dir)
    OUT_DIR = os.path.abspath(output_dir)
    for config in FILES.values():
        config["path"] = os.path.join(BASE, config["filename"])


def rows_from_sheet(ws):
    rows = list(ws.iter_rows(values_only=True))
    headers = [str(h) if h is not None else "" for h in rows[0]]
    out = []
    for values in rows[1:]:
        row = {}
        for index, header in enumerate(headers):
            value = values[index] if index < len(values) else ""
            row[header] = "" if value is None else value
        out.append(row)
    return out


def int_or_blank(value):
    if value == "" or value is None:
        return ""
    try:
        return int(value)
    except (TypeError, ValueError):
        return value


def read_model(model_key, config):
    wb = openpyxl.load_workbook(config["path"], read_only=True, data_only=True)
    summary = {row["metric"]: row for row in rows_from_sheet(wb["Summary"])}
    ranks = rows_from_sheet(wb[config["rank_sheet"]])
    evidence = rows_from_sheet(wb[config["evidence_sheet"]])
    evidence_by_catalog = {str(row["catalog_id"]): row for row in evidence}
    ranks_by_catalog = {str(row["catalog_id"]): row for row in ranks}
    return {
        "summary": summary,
        "ranks": ranks,
        "evidence": evidence,
        "evidence_by_catalog": evidence_by_catalog,
        "ranks_by_catalog": ranks_by_catalog,
    }


def metric(model_data, name):
    row = model_data["summary"].get(name, {})
    return row.get("value", "")


def rank_delta(old_rank, new_rank):
    if old_rank == "" or new_rank == "":
        return ""
    return int(old_rank) - int(new_rank)


def movement_label(delta):
    if delta == "":
        return ""
    if delta > 0:
        return "Moved older/higher"
    if delta < 0:
        return "Moved younger/lower"
    return "No rank change"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, size=8.2, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.0)
            cells[i].width = widths[i]
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9.4)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_paragraph(doc, text, size=9.4, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    return p


def pdf_table(data, widths, header=True, font_size=7.2):
    header_style = ParagraphStyle(
        "TableHeader",
        fontName="Helvetica-Bold",
        fontSize=font_size,
        leading=font_size + 1.2,
        textColor=colors.white,
    )
    cell_style = ParagraphStyle(
        "TableCell",
        fontName="Helvetica",
        fontSize=font_size,
        leading=font_size + 1.4,
        textColor=colors.black,
    )
    wrapped = []
    for row_index, row in enumerate(data):
        wrapped_row = []
        for cell in row:
            style = header_style if header and row_index == 0 else cell_style
            text = str(cell).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            wrapped_row.append(Paragraph(text, style))
        wrapped.append(wrapped_row)
    table = Table(wrapped, colWidths=widths, repeatRows=1 if header else 0)
    style = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9E2EC")),
    ]
    if header:
        style.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ])
    table.setStyle(TableStyle(style))
    return table


def build_pdf(pdf_path, info):
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BriefTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=18,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=4,
    )
    h_style = ParagraphStyle(
        "BriefHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=12,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=6,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "BriefBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.6,
        leading=10.3,
        spaceAfter=4,
    )
    small_style = ParagraphStyle(
        "BriefSmall",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=7.5,
        leading=9,
        textColor=colors.HexColor("#555555"),
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "BriefBullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-6,
        bulletIndent=0,
        spaceAfter=2,
    )

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )
    story = []
    story.append(Paragraph("Kona Biopsy Age Ranking Models: Comparison Brief", title_style))
    story.append(Paragraph("Prepared from attached Models 1-3 workbooks. No database records or raw data were modified for this comparison.", small_style))
    story.append(Paragraph("Bottom Line", h_style))
    story.append(Paragraph("The age-ranking system is a pairwise evidence model, not an absolute age estimate. A sample ranks older when there is dated evidence that its manta was mature before another sampled manta was documented as immature. Tied ranks are expected when the evidence cannot distinguish samples.", body_style))
    story.append(Paragraph(f"The major shift occurred from Model 1 to Model 2: {info['moved_up_12']} samples moved older/higher, {info['moved_down_12']} moved younger/lower, and {info['same_12']} retained the same rank. Model 3 did not change any ranks or scores relative to Model 2 because Model 2 was already treating pup labels as immature; Model 3 made the pup-as-birth-year assumption explicit and added audit columns.", body_style))

    story.append(Paragraph("What Each Model Includes", h_style))
    story.append(Paragraph("Each model keeps the same 83 biopsy samples. The difference is not the sample set; it is which dated maturity/immaturity evidence is allowed to create pairwise older-than comparisons. The models are nested: Model 2 includes everything in Model 1, and Model 3 includes everything in Model 2.", body_style))
    model_data = [["Model", "Explicitly includes", "New vs prior", "Still excludes / caveat"]]
    model_data.extend([[str(cell) for cell in row] for row in info["model_rows"]])
    story.append(pdf_table(model_data, [0.7 * inch, 2.15 * inch, 1.65 * inch, 2.25 * inch], font_size=6.5))
    story.append(Paragraph("Size treatment: raw size measurements were not re-modeled independently in any of these three models. However, HAMER age-class calls may already include size-based maturity or immaturity decisions made in the field/database, so that size-informed evidence is incorporated through the HAMER age-class field in Model 1 and carried forward into Models 2 and 3.", body_style))
    story.append(Paragraph(f"Evidence density increases from {info['model1_edges']} pairwise edges in Model 1 to {info['model2_edges']} in Model 2. Model 3 remains at {info['model3_edges']} edges because, in these files, the new pup-birth-year fields clarify the evidence rather than adding new rank-separating comparisons.", body_style))
    story.append(Spacer(1, 5))

    story.append(Paragraph("How Rankings Changed", h_style))
    movement_data = [["Transition", "Moved older", "Moved younger", "Same rank", "Score changed"]]
    movement_data.extend([[str(cell) for cell in row] for row in info["movement_rows"]])
    story.append(pdf_table(movement_data, [1.35 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch, 1.1 * inch], font_size=7.4))
    story.append(Paragraph("Interpretation: Model 2 supplies many more immature anchors through MPRF age classes, increasing the number of pairwise comparisons and spreading the ranking. Older/mature individuals tend to move higher because they are now ordered before many more MPRF juvenile/pup samples. Samples with only immature evidence tend to move lower because they are now explicitly younger than more mature-dated samples.", body_style))

    story.append(PageBreak())
    story.append(Paragraph("Largest Model 1 to Model 2 Rank Movements", h_style))
    mover_rows = [["Sample", "Source", "M1", "M2", "Delta", "M2 evidence"]]
    for row in info["top_up"][:4]:
        mover_rows.append([row["name"], row["source_group"], row["model1_rank"], row["model2_rank"], row["rank_change_model1_to_model2"], row["model2_evidence"]])
    for row in info["top_down"][:4]:
        mover_rows.append([row["name"], row["source_group"], row["model1_rank"], row["model2_rank"], row["rank_change_model1_to_model2"], row["model2_evidence"]])
    story.append(pdf_table(mover_rows, [2.3 * inch, 0.65 * inch, 0.45 * inch, 0.45 * inch, 0.5 * inch, 1.2 * inch], font_size=7.2))
    story.append(Spacer(1, 5))

    story.append(Paragraph("Confidence Context", h_style))
    bullets = [
        "Model 1 is the cleanest parent/offspring exclusion screen because it relies on HAMER field maturity calls, which may already incorporate HAMER size thresholds, and long MPRF resight histories rather than broad MPRF age-class assumptions.",
        "Model 2 is useful for ranking all 83 samples but lower confidence for MPRF-only age classes, because the underlying maturity criteria may not be consistently documented.",
        "Model 3 is not a new rank outcome in the current files; its value is transparency. It preserves the pup birth-year anchor and derived +15-year maturity dates for review.",
        "All models remain sensitive to ongoing QC of sighting, manta, photo, biopsy, and catalog links. Treat ranks as relative evidence weights, not definitive ages.",
    ]
    for item in bullets:
        story.append(Paragraph(f"• {item}", bullet_style))

    story.append(Paragraph("Recommended Next Checks", h_style))
    checks = [
        "Review the 32 pup birth-year anchors and confirm whether each label is defensible as a near-birth-year observation.",
        "For MPRF adult/juvenile labels, inspect notes, photos, size records, mating scars, pregnancy, mating-train context, and clasper evidence where available.",
        "Add independent raw-size modeling only after checking measurement consistency over time; large up/down size fluctuations should be flagged before use.",
        "Use the master CSV as the audit table: it shows every sample, all model ranks/scores, evidence basis, movement direction, and Model 3 pup/+15-year fields.",
    ]
    for item in checks:
        story.append(Paragraph(f"• {item}", bullet_style))
    story.append(Paragraph("Source: attached Model 1, Model 2, and Model 3 workbooks in the Kona Biopsy Age Rankings folder.", small_style))
    doc.build(story)


def build():
    os.makedirs(OUT_DIR, exist_ok=True)
    models = {key: read_model(key, config) for key, config in FILES.items()}

    master_rows = []
    for catalog_id, model3_rank in sorted(
        models["model3"]["ranks_by_catalog"].items(),
        key=lambda item: (int(item[1]["model3_rank"]), item[1]["name"]),
    ):
        m1 = models["model1"]["ranks_by_catalog"][catalog_id]
        m2 = models["model2"]["ranks_by_catalog"][catalog_id]
        m3 = models["model3"]["ranks_by_catalog"][catalog_id]
        e1 = models["model1"]["evidence_by_catalog"].get(catalog_id, {})
        e2 = models["model2"]["evidence_by_catalog"].get(catalog_id, {})
        e3 = models["model3"]["evidence_by_catalog"].get(catalog_id, {})
        d12 = rank_delta(m1["model1_rank"], m2["model2_rank"])
        d23 = rank_delta(m2["model2_rank"], m3["model3_rank"])
        d13 = rank_delta(m1["model1_rank"], m3["model3_rank"])
        master_rows.append({
            "catalog_id": catalog_id,
            "name": m3["name"],
            "biopsy_id": m3["biopsy_id"],
            "sample_id": m3["sample_id"],
            "sample_lookup_name": m3["sample_lookup_name"],
            "source_group": m3["source_group"],
            "model1_rank": int_or_blank(m1["model1_rank"]),
            "model1_score": int_or_blank(m1["model1_score"]),
            "model1_evidence": m1["model1_evidence"],
            "model1_mature_date": m1["mature_date"],
            "model1_immature_date": m1["immature_date"],
            "model1_older_than_count": int_or_blank(m1["older_than_count"]),
            "model1_younger_than_count": int_or_blank(m1["younger_than_count"]),
            "model1_basis": e1.get("model1_evidence_basis", ""),
            "model2_rank": int_or_blank(m2["model2_rank"]),
            "model2_score": int_or_blank(m2["model2_score"]),
            "model2_evidence": m2["model2_evidence"],
            "model2_mature_date": m2["mature_date"],
            "model2_immature_date": m2["immature_date"],
            "model2_older_than_count": int_or_blank(m2["older_than_count"]),
            "model2_younger_than_count": int_or_blank(m2["younger_than_count"]),
            "model2_basis": e2.get("model2_evidence_basis", ""),
            "model3_rank": int_or_blank(m3["model3_rank"]),
            "model3_score": int_or_blank(m3["model3_score"]),
            "model3_evidence": m3["model3_evidence"],
            "model3_mature_date": m3["mature_date"],
            "model3_immature_date": m3["immature_date"],
            "model3_older_than_count": int_or_blank(m3["older_than_count"]),
            "model3_younger_than_count": int_or_blank(m3["younger_than_count"]),
            "model3_basis": e3.get("model3_evidence_basis", ""),
            "model3_birth_year_anchor_date": e3.get("birth_year_anchor_date", ""),
            "model3_immature_plus_15_mature_date": e3.get("immature_plus_15_mature_date", ""),
            "model3_last_evidence_date": e3.get("last_evidence_date", ""),
            "rank_change_model1_to_model2": d12,
            "movement_model1_to_model2": movement_label(d12),
            "rank_change_model2_to_model3": d23,
            "movement_model2_to_model3": movement_label(d23),
            "rank_change_model1_to_model3": d13,
            "movement_model1_to_model3": movement_label(d13),
            "current_view_rank": m3.get("current_view_rank", ""),
            "rank_note": m3.get("rank_note", ""),
        })

    csv_path = os.path.join(OUT_DIR, "kona_age_ranking_models_master_comparison.csv")
    with open(csv_path, "w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(master_rows[0].keys()))
        writer.writeheader()
        writer.writerows(master_rows)

    deltas_12 = [row["rank_change_model1_to_model2"] for row in master_rows]
    deltas_23 = [row["rank_change_model2_to_model3"] for row in master_rows]
    moved_up_12 = sum(1 for d in deltas_12 if d != "" and d > 0)
    moved_down_12 = sum(1 for d in deltas_12 if d != "" and d < 0)
    same_12 = sum(1 for d in deltas_12 if d == 0)
    moved_up_23 = sum(1 for d in deltas_23 if d != "" and d > 0)
    moved_down_23 = sum(1 for d in deltas_23 if d != "" and d < 0)
    same_23 = sum(1 for d in deltas_23 if d == 0)

    top_up = sorted(master_rows, key=lambda row: row["rank_change_model1_to_model2"], reverse=True)[:6]
    top_down = sorted(master_rows, key=lambda row: row["rank_change_model1_to_model2"])[:6]
    score_changed_12 = sum(1 for row in master_rows if row["model1_score"] != row["model2_score"])
    score_changed_23 = sum(1 for row in master_rows if row["model2_score"] != row["model3_score"])
    unresolved_counts = {
        "Model 1": metric(models["model1"], "Samples without Model 1 evidence") or metric(models["model1"], "Samples unresolved by Model 1") or 47,
        "Model 2": metric(models["model2"], "Samples unresolved by Model 2") or metric(models["model2"], "Samples without Model 2 evidence") or 8,
        "Model 3": metric(models["model3"], "Samples unresolved by Model 3") or 8,
    }

    docx_path = os.path.join(OUT_DIR, "kona_age_ranking_model_comparison_brief.docx")
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.4)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(9.4)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Kona Biopsy Age Ranking Models: Comparison Brief")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run(f"Prepared from attached Models 1-3 workbooks | {datetime.now().strftime('%Y-%m-%d %H:%M %Z')}")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "Bottom Line")
    add_paragraph(
        doc,
        "The age-ranking system is a pairwise evidence model, not an absolute age estimate. A sample ranks older when there is dated evidence that its manta was mature before another sampled manta was documented as immature. Tied ranks are expected when the evidence cannot distinguish samples."
    )
    add_paragraph(
        doc,
        f"The major shift occurred from Model 1 to Model 2: {moved_up_12} samples moved older/higher, {moved_down_12} moved younger/lower, and {same_12} retained the same rank. Model 3 did not change any ranks or scores relative to Model 2 because Model 2 was already treating pup labels as immature; Model 3 made the pup-as-birth-year assumption explicit and added audit columns."
    )

    add_heading(doc, "What Each Model Includes")
    add_paragraph(
        doc,
        "Each model keeps the same 83 biopsy samples. The difference is not the sample set; it is which dated maturity/immaturity evidence is allowed to create pairwise older-than comparisons. The models are nested: Model 2 includes everything in Model 1, and Model 3 includes everything in Model 2."
    )
    model_table_rows = [
        [
            "Model 1",
            "HAMER biopsy-linked age class, which may already include size-based maturity/immaturity decisions, plus MPRF >15-year resight history counted as mature at first sighting +15 years.",
            "Starting model.",
            "Does not use MPRF adult/juvenile/pup labels, pup-as-birth-year, independent raw-size modeling, notes, photos, or sex-specific evidence unless already embedded in the HAMER age class.",
        ],
        [
            "Model 2",
            "Everything in Model 1, plus MPRF-provided age classes. Adult/mature is treated as mature; juvenile/pup/immature is treated as immature on the MPRF first/effective sighting date.",
            "Adds medium-confidence MPRF age-class labels, increasing pairwise edges from 43 to 597.",
            "Still does not separate pup from juvenile as a birth-year anchor, and does not independently verify the field criteria behind MPRF age classes.",
        ],
        [
            "Model 3",
            "Everything in Model 2, plus explicit pup birth-year anchors and immature/pup +15-year mature-date fields when later sighting history supports that derived date.",
            "Adds transparency/audit fields for pup and +15-year logic; no rank or score changes in the current files.",
            "Still does not independently re-evaluate raw size measurements, measurement trends, reproductive notes, clasper notes, photos, or unresolved QC corrections.",
        ],
    ]
    add_table(
        doc,
        ["Model", "Explicitly includes", "New vs prior", "Still excludes / caveat"],
        model_table_rows,
        [Inches(0.75), Inches(2.35), Inches(1.55), Inches(2.15)],
    )
    add_paragraph(
        doc,
        "Size treatment: raw size measurements were not re-modeled independently in any of these three models. However, HAMER age-class calls may already include size-based maturity or immaturity decisions made in the field/database, so that size-informed evidence is incorporated through the HAMER age-class field in Model 1 and carried forward into Models 2 and 3."
    )
    add_paragraph(
        doc,
        f"Evidence density increases from {metric(models['model1'], 'Pairwise older-than edges') or 43} pairwise edges in Model 1 to {metric(models['model2'], 'Pairwise older-than edges') or 597} in Model 2. Model 3 remains at {metric(models['model3'], 'Pairwise older-than edges') or 597} edges because, in these files, the new pup-birth-year fields clarify the evidence rather than adding new rank-separating comparisons."
    )

    add_heading(doc, "How Rankings Changed")
    movement_rows = [
        ["Model 1 -> 2", moved_up_12, moved_down_12, same_12, score_changed_12],
        ["Model 2 -> 3", moved_up_23, moved_down_23, same_23, score_changed_23],
    ]
    add_table(
        doc,
        ["Transition", "Moved older", "Moved younger", "Same rank", "Score changed"],
        movement_rows,
        [Inches(1.4), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.1)],
    )
    add_paragraph(
        doc,
        "Interpretation: Model 2 supplies many more immature anchors through MPRF age classes, which increases the number of pairwise comparisons and spreads the ranking. Older/mature individuals tend to move higher because they are now ordered before many more MPRF juvenile/pup samples. Samples with only immature evidence tend to move lower because they are now explicitly younger than more mature-dated samples."
    )

    doc.add_page_break()
    add_heading(doc, "Largest Model 1 to Model 2 Rank Movements")
    mover_rows = []
    for row in top_up[:4]:
        mover_rows.append([row["name"], row["source_group"], row["model1_rank"], row["model2_rank"], row["rank_change_model1_to_model2"], row["model2_evidence"]])
    for row in top_down[:4]:
        mover_rows.append([row["name"], row["source_group"], row["model1_rank"], row["model2_rank"], row["rank_change_model1_to_model2"], row["model2_evidence"]])
    add_table(
        doc,
        ["Sample", "Source", "M1 rank", "M2 rank", "Delta", "M2 evidence"],
        mover_rows,
        [Inches(2.1), Inches(0.7), Inches(0.65), Inches(0.65), Inches(0.55), Inches(1.35)],
    )

    add_heading(doc, "Confidence Context")
    add_bullet(doc, "Model 1 is the cleanest screen for parent/offspring exclusion because it relies on HAMER field maturity calls, which may already incorporate HAMER size thresholds, and long MPRF resight histories rather than broad MPRF age-class assumptions.")
    add_bullet(doc, "Model 2 is useful for ranking all 83 samples but lower confidence for MPRF-only age classes, because the underlying maturity criteria may not be consistently documented.")
    add_bullet(doc, "Model 3 is not a new rank outcome in the current files; its value is transparency. It preserves the pup birth-year anchor and the derived +15-year maturity dates for review.")
    add_bullet(doc, "All models remain sensitive to ongoing QC of sighting, manta, photo, biopsy, and catalog links. Treat ranks as relative evidence weights, not definitive ages.")

    add_heading(doc, "Recommended Next Checks")
    add_bullet(doc, "Review the 32 pup birth-year anchors and confirm whether each pup label is defensible as a near-birth-year observation.")
    add_bullet(doc, "For MPRF adult/juvenile labels, inspect notes, photos, size records, mating scars, pregnancy, mating-train context, and clasper evidence where available.")
    add_bullet(doc, "Add independent raw-size modeling only after checking measurement consistency over time; large up/down size fluctuations should be flagged before being used as maturity evidence.")
    add_bullet(doc, "Keep the master CSV as the audit table: it shows every sample, all model ranks/scores, evidence basis, movement direction, and Model 3 pup/+15-year fields.")

    source_note = doc.add_paragraph()
    source_note.paragraph_format.space_before = Pt(6)
    source_note.paragraph_format.space_after = Pt(0)
    run = source_note.add_run(
        "Source: attached Model 1, Model 2, and Model 3 workbooks in the Kona Biopsy Age Rankings folder. No database records or raw data were modified for this comparison."
    )
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(docx_path)

    pdf_path = os.path.join(OUT_DIR, "kona_age_ranking_model_comparison_brief.pdf")
    build_pdf(
        pdf_path,
        {
            "moved_up_12": moved_up_12,
            "moved_down_12": moved_down_12,
            "same_12": same_12,
            "moved_up_23": moved_up_23,
            "moved_down_23": moved_down_23,
            "same_23": same_23,
            "score_changed_12": score_changed_12,
            "score_changed_23": score_changed_23,
            "top_up": top_up,
            "top_down": top_down,
            "model_rows": model_table_rows,
            "movement_rows": movement_rows,
            "model1_edges": metric(models["model1"], "Pairwise older-than edges") or 43,
            "model2_edges": metric(models["model2"], "Pairwise older-than edges") or 597,
            "model3_edges": metric(models["model3"], "Pairwise older-than edges") or 597,
        },
    )

    return csv_path, docx_path, pdf_path, {
        "total_rows": len(master_rows),
        "moved_up_12": moved_up_12,
        "moved_down_12": moved_down_12,
        "same_12": same_12,
        "moved_up_23": moved_up_23,
        "moved_down_23": moved_down_23,
        "same_23": same_23,
        "score_changed_12": score_changed_12,
        "score_changed_23": score_changed_23,
        "model1_edges": metric(models["model1"], "Pairwise older-than edges") or 43,
        "model2_edges": metric(models["model2"], "Pairwise older-than edges") or 597,
        "model3_edges": metric(models["model3"], "Pairwise older-than edges") or 597,
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Recreate the historical May 29, 2026 Kona Model 1/2/3 comparison outputs."
    )
    parser.add_argument("--input-dir", required=True, help="Directory containing the three historical model workbooks.")
    parser.add_argument("--output-dir", required=True, help="Directory where the historical outputs should be written.")
    args = parser.parse_args()
    configure_paths(args.input_dir, args.output_dir)
    csv_output, docx_output, pdf_output, stats = build()
    print({"csv": csv_output, "docx": docx_output, "pdf": pdf_output, "stats": stats})
