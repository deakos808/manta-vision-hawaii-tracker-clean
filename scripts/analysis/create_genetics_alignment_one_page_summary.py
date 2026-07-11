from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("reports/kona_pi_hat_generation_alignment")
OUT_DOCX = OUT_DIR / "Kona_PI_HAT_PCSU_Maturity_Sensitivity_One_Page_Summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=8.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "C9D3DF")


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_body(doc, text, size=9.4, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("• ")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.bold = True
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Kona PI_HAT Alignment with PCSU Generation Matrix")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(5)
    r = subtitle.add_run("One-page interpretation of strong genetic pairs and maturity-age sensitivity")
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor(84, 99, 119)

    add_heading(doc, "Bottom Line")
    add_body(
        doc,
        "Using PI_HAT >= 0.35 as the strong-relatedness screen, 27 genetic pairs were flagged. "
        "Nineteen mapped cleanly to the Kona biopsied PCSU matrix. Of those mapped strong pairs, four were compatible with a directional parent-child hypothesis, one was compatible with same-generation/sibling plausibility, and fourteen remained age-unresolved (U/U).",
    )
    add_body(
        doc,
        "The five biologically resolved strong pairs were unchanged when maturity-age thresholds were rerun at the low, midpoint, and high settings: males 5/6.5/8 years and females 8/11.5/15 years. This is a useful robustness result.",
    )

    add_heading(doc, "Resolved Strong PI_HAT Pairs")
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Inches(1.55), Inches(1.55), Inches(0.75), Inches(1.1), Inches(2.25)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
    headers = ["Sample A", "Sample B", "PI_HAT", "PCSU call", "Interpretation"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")

    pair_rows = [
        ("BI105 Big Bertha", "BI107 Conae", "0.531", "P/C", "Big Bertha could be parent of Conae"),
        ("BI105 Big Bertha", "BI123 Koen", "0.489", "P/C", "Big Bertha could be parent of Koen"),
        ("BI136 Takahashi", "BI123 Koen", "0.471", "P/C", "Takahashi could be parent of Koen"),
        ("BI100 Koie", "BI131 Black Diamond", "0.444", "P/C", "Koie could be parent of Black Diamond"),
        ("BI107 Conae", "BI_K05 Maluhia", "0.503", "S/S", "Same-generation / sibling-compatible under age evidence"),
    ]
    for row_values in pair_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value, size=7.8)
    set_table_borders(table)

    add_heading(doc, "Maturity-Age Sensitivity")
    add_bullet(doc, "Full Kona biopsied matrix: low maturity ages increased P-C relationships from 253 to 357; high maturity ages reduced them to 192 and increased S relationships from 529 to 590.")
    add_bullet(doc, "Strong PI_HAT subset: P-C and S classifications did not change across low, midpoint, or high maturity-age thresholds.")
    add_bullet(doc, "This is because the resolved strong pairs already have enough dated life-stage separation to stay in the same category, or they lack the specific dated adult-versus-juvenile evidence needed to move out of U/U.")

    add_heading(doc, "Ranking Versus Generation-Gap Logic")
    add_body(
        doc,
        "The low/mid/high reruns recalculated both the age/ranking model and the generation-gap PCSU matrix. "
        "Therefore, the stability of the five resolved strong PI_HAT calls is not caused by holding midpoint ages fixed for ranking while changing only the matrix threshold.",
        size=9.1,
    )
    add_body(
        doc,
        "For Kona biopsies, most minimum ages are anchored by first sighting dates, so maturity-age assumptions have limited effect on rank order. "
        "They have a larger effect on broad pairwise generation counts, but the genetically strongest resolved pairs were robust because their PCSU classification was supported by observed timing rather than a fragile rank shift.",
        size=9.1,
        after=2,
    )

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(2)
    foot.paragraph_format.space_after = Pt(0)
    rf = foot.add_run(
        "Interpretation note: PI_HAT identifies close genetic similarity; PCSU is a life-history plausibility filter. P/C or S/S calls are compatibility statements, not standalone proof of parentage."
    )
    rf.font.name = "Calibri"
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(84, 99, 119)
    rf.italic = True

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
