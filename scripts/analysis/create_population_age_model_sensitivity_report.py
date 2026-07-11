from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("reports/population_age_model_sensitivity")
DATA_PATH = OUT_DIR / "population_age_model_sensitivity_report_data.json"
DOCX_PATH = OUT_DIR / "Population_Age_Model_Sensitivity_Comparison_Report.docx"
SUPPLEMENTAL_WORKBOOK = OUT_DIR / "population_age_model_sensitivity_tables.xlsx"


def fmt(value, suffix: str = "") -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        if abs(value - round(value)) < 0.05:
            return f"{value:.0f}{suffix}"
        return f"{value:.1f}{suffix}"
    return f"{value}{suffix}"


def pct(numerator: float, denominator: float) -> str:
    if not denominator:
        return "-"
    return f"{100 * numerator / denominator:.1f}%"


def scope(data: dict, key: str) -> dict:
    return next(item for item in data["scopes"] if item["key"] == key)


def model(scope_data: dict, key: str) -> dict | None:
    return next((item for item in scope_data["summaries"] if item["key"] == key), None)


def sensitivity(scope_data: dict, model_key: str, variant: str) -> dict | None:
    return next(
        (
            item
            for item in scope_data.get("maturityAgeSensitivity", [])
            if item["modelKey"] == model_key and item["maturityVariant"] == variant
        ),
        None,
    )


def short_model(label: str) -> str:
    return (
        label.replace("Model 1 + ", "")
        .replace("Best evidence model: ", "Best evidence: ")
        .replace(" with maturity-size/age assumptions", "")
        .replace(" with maturity-age assumptions", "")
        .replace("HAMER dated age class", "HAMER age class")
        .replace("HAMER size evidence", "HAMER size")
        .replace("MPRF dated/first-sighting age class", "MPRF age class")
    )


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)
    return doc


def build_scope_rows(data: dict) -> list[list[object]]:
    rows = []
    for key in ["kona_biopsied", "kona_all", "maui_nui_biopsied", "maui_nui_all"]:
        s = scope(data, key)
        m1 = model(s, "m1")
        best = model(s, "best_evidence_no_mprf_class")
        all_models = model(s, "all_models")
        rows.append(
            [
                s["label"],
                s["recordCount"],
                best["minAge"]["changedCount"],
                best["ranks"]["relativeOrderChangedCount"],
                fmt(m1["minAge"]["mean"]),
                fmt(best["minAge"]["mean"]),
                best["pairwise"]["animalsWithP"],
                best["pairwise"]["animalsWithC"],
                best["pairwise"]["pcRelationships"],
                best["pairwise"]["sameGenerationRelationships"],
                all_models["minAge"]["changedCount"] if all_models else "n/a",
                all_models["ranks"]["relativeOrderChangedCount"] if all_models else "n/a",
            ]
        )
    return rows


def build_population_summary_rows(data: dict) -> list[list[object]]:
    rows = []
    for key in ["kona_biopsied", "kona_all", "maui_nui_biopsied", "maui_nui_all"]:
        s = scope(data, key)
        p = s["populationSummary"]
        rows.append(
            [
                s["label"],
                p["assessableMantas"],
                p["uniqueBiopsiedMantas"],
                p["biopsyRecords"],
                p["duplicateBiopsyRecordsCollapsed"],
                p["withFirstSighting"],
                p["pupAtFirstSighting"],
                p["withSizeEvidence"],
                p["withHamerAgeClass"],
                p["withMprfAgeClass"],
                p["adultAgeClassEvidence"],
                p["juvenileAgeClassEvidence"],
                p["male"],
                p["female"],
                p["unknownSex"],
            ]
        )
    return rows


def build_model_rows(scope_data: dict) -> list[list[object]]:
    preferred_order = [
        "m1",
        "m2_pup",
        "m3_size_assumption",
        "m4_hamer_age_class_assumption",
        "best_evidence_no_mprf_class",
        "m5_mprf_age_class_assumption",
        "all_models",
    ]
    rows = []
    for key in preferred_order:
        m = model(scope_data, key)
        if not m:
            continue
        rows.append(
            [
                short_model(m["label"]),
                m["minAge"]["changedCount"],
                fmt(m["minAge"]["meanDeltaChanged"]),
                m["ranks"]["relativeOrderChangedCount"],
                fmt(m["ranks"]["meanAbsRankDelta"]),
                m["pairwise"]["animalsWithP"],
                m["pairwise"]["animalsWithC"],
                m["pairwise"]["pcRelationships"],
                m["pairwise"]["sameGenerationRelationships"],
            ]
        )
    return rows


def build_sensitivity_rows(scope_data: dict) -> list[list[object]]:
    rows = []
    for variant in ["low", "midpoint", "high"]:
        item = sensitivity(scope_data, "best_evidence_no_mprf_class", variant)
        if not item:
            continue
        rows.append(
            [
                item["maturityLabel"],
                f"{fmt(item['maleMaturityAgeYears'])}/{fmt(item['femaleMaturityAgeYears'])}",
                fmt(item["meanMinimumAge"]),
                item["relativeOrderChangedVsMidpoint"],
                item["mantasWithP"],
                item["mantasWithC"],
                item["pcRelationships"],
                item["sameGenerationRelationships"],
            ]
        )
    return rows


def top_parents_text(scope_data: dict, n: int = 6) -> str:
    top = scope_data.get("topParentContributors", [])[:n]
    if not top:
        return "No parent-position contributors were available."
    return "; ".join(f"{item['name']} ({item['P']} P calls)" for item in top)


def top_age_change_text(scope_data: dict, model_key: str = "best_evidence_no_mprf_class", n: int = 5) -> str:
    m = model(scope_data, model_key)
    rows = m.get("topAgeIncreases", [])[:n] if m else []
    if not rows:
        return "No individual minimum-age increases were detected."
    return "; ".join(f"{r['name']} (+{fmt(r['ageIncrease'])} yrs)" for r in rows)


def build_doc() -> None:
    data = json.loads(DATA_PATH.read_text())
    doc = setup_document()
    kona_b = scope(data, "kona_biopsied")
    kona_all = scope(data, "kona_all")
    maui_b = scope(data, "maui_nui_biopsied")
    maui_all = scope(data, "maui_nui_all")
    kona_b_best = model(kona_b, "best_evidence_no_mprf_class")
    kona_all_best = model(kona_all, "best_evidence_no_mprf_class")
    maui_b_best = model(maui_b, "best_evidence_no_mprf_class")
    maui_all_best = model(maui_all, "best_evidence_no_mprf_class")

    title = doc.add_heading("Population Age Model Sensitivity Comparison", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Kona and Maui Nui manta age intervals, ranks, and pairwise generation diagnostics; "
        f"age estimates as of {data['ageReferenceDate']}."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "This report extends the Kona biopsy age-ranking sensitivity analysis to four active scopes: Kona biopsied mantas, "
        "Kona all assessable catalog mantas, Maui Nui biopsied mantas, and Maui Nui all assessable catalog mantas. Full-population "
        "datasets include catalog mantas with at least one dated sighting; catalog records without a dated sighting are excluded. "
        "Maui Nui is treated as a single population including Maui, Molokai, Lanai, and Kahoolawe. MPRF age-class evidence is only "
        "evaluated for Kona."
    )
    add_bullets(
        doc,
        [
            f"Kona remains highly stable: the best evidence model (Models 1-4, excluding MPRF age class) changes only {kona_b_best['minAge']['changedCount']} of {kona_b['recordCount']} Kona biopsy ages and {kona_all_best['minAge']['changedCount']} of {kona_all['recordCount']} Kona all-manta ages.",
            f"Maui Nui is much more sensitive to HAMER size and HAMER age-class evidence: the best evidence model changes {maui_b_best['minAge']['changedCount']} of {maui_b['recordCount']} unique Maui Nui biopsy-manta ages and {maui_all_best['minAge']['changedCount']} of {maui_all['recordCount']} Maui Nui all-manta ages.",
            f"The Maui Nui database query returned {maui_b['populationSummary']['biopsyRecords']} biopsy/sample records in the Maui Nui scope, but these collapse to {maui_b['populationSummary']['uniqueBiopsiedMantas']} unique biopsied mantas after repeated catalog IDs are deduplicated.",
            "Pup-at-first-sighting evidence rarely changes minimum age or rank because first sighting already supplies the minimum-age floor, but it improves same-generation resolution when pup labels are available.",
            "Pairwise generation diagnostics are more sensitive to maturity-age thresholds than rank order is, but P/C calls are now based only on dated adult-versus-juvenile evidence. Lower maturity ages increase P-C calls; higher maturity ages reduce P-C calls and push more pairs toward same-generation or unresolved classifications.",
            "The best default model remains a HAMER-prioritized evidence model: first sighting, pup-at-first-sighting, HAMER size evidence, and HAMER dated age class. MPRF age class should remain a Kona-only sensitivity layer until its dated labels are validated.",
        ],
    )

    add_heading(doc, "Methods", 1)
    add_body(
        doc,
        "Five evidence models were evaluated against Model 1, the first-sighting baseline. Model 1 uses only the earliest known "
        "dated sighting and is treated as the least ambiguous baseline. Model 2 adds pup-at-first-sighting labels as direct "
        "birth-anchor evidence. Model 3 adds HAMER size measurements interpreted using sex-specific maturity-size and maturity-age "
        "assumptions. Model 4 adds HAMER dated age-class observations interpreted with sex-specific maturity ages. Model 5 adds MPRF "
        "dated or first-sighting age-class observations and is only evaluated for Kona. The recommended best evidence model combines "
        "Models 1-4 and excludes MPRF age class."
    )
    add_body(
        doc,
        "Default maturity ages were set to 6.5 years for males and 11.5 years for females, with sensitivity tests at low values "
        "(male 5, female 8) and high values (male 8, female 15). Default maturity sizes were 2.8 m for males and 3.37 m for females. "
        "Growth-rate extrapolation was not used for primary ranking or generation diagnostics because existing growth evidence is too "
        "sparse and stage-dependent for reliable linear age assignment."
    )
    add_body(
        doc,
        "Pairwise generation diagnostics classify ordered pairs as P, C, S, or U. P indicates that the row animal was observed adult/mature "
        "on a dated record and the column animal was observed later as juvenile/pup or below maturity size after enough elapsed time to clear "
        "the possible child's sex-specific maturity threshold. C is the reciprocal child-position state. S indicates dated stage evidence exists "
        "but the separation is shorter than the relevant maturity threshold. U indicates insufficient dated stage evidence. Back-calculated "
        "minimum ages are used for rank ordering, but they are not used to back-date adult status for P/C generation calls."
    )

    add_heading(doc, "Cross-scope Summary", 1)
    add_heading(doc, "Population Evidence Summary", 2)
    add_table(
        doc,
        [
            "Scope",
            "Assessable mantas",
            "Unique biopsied mantas",
            "Biopsy records",
            "Collapsed duplicate biopsy records",
            "First sighting",
            "Pup first sighting",
            "Size evidence",
            "HAMER age class",
            "MPRF age class",
            "Adult evidence",
            "Juvenile evidence",
            "Males",
            "Females",
            "Unknown sex",
        ],
        build_population_summary_rows(data),
    )
    add_body(
        doc,
        "Counts are manta-level counts after deduplication by HAMER catalog/manta identity. The biopsy-record column is retained because "
        "repeat biopsy/sample rows can exceed the number of unique animals, especially in Maui Nui. Pairwise matrices and rank summaries use "
        "the unique-manta count, not raw biopsy-record count."
    )

    add_heading(doc, "Model Outcome Summary", 2)
    add_table(
        doc,
        [
            "Scope",
            "Records",
            "Best-model age changes",
            "Best-model relative rank changes",
            "M1 mean age",
            "Best mean age",
            "Mantas P",
            "Mantas C",
            "P-C calls",
            "S calls",
            "All-model age changes",
            "All-model relative rank changes",
        ],
        build_scope_rows(data),
    )
    add_body(
        doc,
        "The contrast between Kona and Maui Nui is the central result. In Kona, additional HAMER evidence is sparse relative to the "
        "long first-sighting history, so the baseline rank structure is already robust. In Maui Nui, HAMER size and dated age-class "
        "records are more influential, so the best evidence model materially increases minimum ages and reshuffles more relative-order "
        "neighbors, especially in the full-population dataset."
    )

    add_heading(doc, "Kona Results", 1)
    add_heading(doc, "Kona Biopsied Mantas", 2)
    add_table(
        doc,
        ["Model", "Age changed", "Mean delta changed", "Relative rank changes", "Mean abs rank delta", "Mantas P", "Mantas C", "P-C calls", "S calls"],
        build_model_rows(kona_b),
    )
    add_body(
        doc,
        f"Kona biopsied animals remain stable under the HAMER-prioritized best model. The main individual age increases are "
        f"{top_age_change_text(kona_b)}. The top parent-position contributors are {top_parents_text(kona_b)}."
    )
    add_heading(doc, "Kona All Assessable Mantas", 2)
    add_table(
        doc,
        ["Model", "Age changed", "Mean delta changed", "Relative rank changes", "Mean abs rank delta", "Mantas P", "Mantas C", "P-C calls", "S calls"],
        build_model_rows(kona_all),
    )
    add_body(
        doc,
        f"Kona all-manta results follow the biopsy pattern: the best model changes only {model(kona_all, 'best_evidence_no_mprf_class')['minAge']['changedCount']} "
        f"of {kona_all['recordCount']} minimum ages. MPRF age class has a larger effect than HAMER-only evidence but remains a sensitivity layer. "
        f"The leading parent-position contributors are {top_parents_text(kona_all)}."
    )

    add_heading(doc, "Maui Nui Results", 1)
    add_heading(doc, "Maui Nui Biopsied Mantas", 2)
    add_table(
        doc,
        ["Model", "Age changed", "Mean delta changed", "Relative rank changes", "Mean abs rank delta", "Mantas P", "Mantas C", "P-C calls", "S calls"],
        build_model_rows(maui_b),
    )
    add_body(
        doc,
        f"Maui Nui biopsied animals are much more influenced by HAMER evidence than Kona biopsied animals. HAMER size evidence changes "
        f"{model(maui_b, 'm3_size_assumption')['minAge']['changedCount']} ages and HAMER age-class evidence changes "
        f"{model(maui_b, 'm4_hamer_age_class_assumption')['minAge']['changedCount']} ages; combined, the best model changes "
        f"{model(maui_b, 'best_evidence_no_mprf_class')['minAge']['changedCount']} of {maui_b['recordCount']} ages. The most affected animals are "
        f"{top_age_change_text(maui_b)}. The top parent-position contributors are {top_parents_text(maui_b)}."
    )
    add_heading(doc, "Maui Nui All Assessable Mantas", 2)
    add_table(
        doc,
        ["Model", "Age changed", "Mean delta changed", "Relative rank changes", "Mean abs rank delta", "Mantas P", "Mantas C", "P-C calls", "S calls"],
        build_model_rows(maui_all),
    )
    add_body(
        doc,
        f"In the full Maui Nui population, the best model changes {model(maui_all, 'best_evidence_no_mprf_class')['minAge']['changedCount']} of "
        f"{maui_all['recordCount']} age estimates and produces {model(maui_all, 'best_evidence_no_mprf_class')['ranks']['relativeOrderChangedCount']} "
        "relative-order changes. This is the strongest evidence that alternative models do not have uniform influence across populations: "
        "Maui Nui has enough dated HAMER evidence to alter many age intervals, whereas Kona's first-sighting history dominates."
    )
    add_body(doc, f"The leading Maui Nui all-manta parent-position contributors are {top_parents_text(maui_all)}.")

    add_heading(doc, "Maturity-age Sensitivity", 1)
    add_body(
        doc,
        "Maturity-age sensitivity confirms the expected biological pattern. Minimum-age ranks are usually stable because only animals with "
        "dated juvenile/adult or size anchors move. Generation diagnostics are more sensitive because the P-C threshold itself changes."
    )
    for s in [kona_b, kona_all, maui_b, maui_all]:
        add_heading(doc, s["label"], 2)
        add_table(
            doc,
            ["Variant", "M/F maturity age", "Mean min age", "Relative rank changes", "Mantas P", "Mantas C", "P-C calls", "S calls"],
            build_sensitivity_rows(s),
        )

    add_heading(doc, "Interpretation and Model Recommendation", 1)
    add_body(
        doc,
        "For Kona, Model 1 alone captures most of the age-ranking structure because long first-sighting history provides strong minimum-age "
        "anchors. Models 2-4 add valuable biological documentation but have small effects on age estimates and ranks. Model 5 is useful for "
        "asking what changes when MPRF labels are allowed, but it should remain a sensitivity model until those date-specific age-class labels "
        "are validated."
    )
    add_body(
        doc,
        "For Maui Nui, the recommended default should still be the HAMER-prioritized best evidence model, but the interpretation is different. "
        "Models 3 and 4 provide substantial additional information because dated size and HAMER age-class records are more common and more "
        "age-informative. This makes Maui Nui the stronger test case for demonstrating why the workbench should support evidence models beyond "
        "first sighting."
    )
    add_body(
        doc,
        "Across all scopes, the practical recommendation is to report both the Model 1 baseline and the best evidence model. The baseline keeps "
        "the analysis transparent and assumption-light. The best evidence model shows what is gained by allowing biologically interpretable "
        "HAMER evidence. Pairwise generation diagnostics should always be presented with maturity-age sensitivity because generation-gap counts "
        "can change substantially even when ranks do not, and because those calls require actual dated adult-versus-juvenile evidence rather than "
        "inferred pre-observation adult dates."
    )

    add_heading(doc, "Supplemental Materials", 1)
    add_body(
        doc,
        f"Supplemental workbook: {SUPPLEMENTAL_WORKBOOK.name}. This workbook contains model summaries and full best-evidence pairwise matrices "
        "for the four scopes. Full-population matrices are intentionally retained in the workbook rather than embedded in this report because "
        "the Maui Nui all-manta matrix includes hundreds of records and is not readable as a manuscript table."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
